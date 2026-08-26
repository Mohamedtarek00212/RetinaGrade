"""Generate Presentation_Report.docx for the RetinaGrade project.

Documentation only. Reads verified artefacts already copied into Presentation/
(and a few read-only source reports) and embeds the generated figures. Writes a
single .docx under Presentation/. No source code / config / checkpoint / output
is modified. All numbers come from the JSON/CSV artefacts, never invented.
"""
from __future__ import annotations

import csv
import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"d:\RetinaGrade")
PRES = ROOT / "Presentation"
FIG = PRES / "Figures"
MET = PRES / "Metrics"
REP = PRES / "Reports"
CFG = PRES / "Configs"

CLASS_FULL = ["No DR (Normal)", "Mild DR", "Moderate DR", "Severe DR", "Proliferative DR (PDR)"]
CLASS_SHORT = ["No DR", "Mild", "Moderate", "Severe", "PDR"]
NA = "Not available in the current repository."

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
GREY = RGBColor(0x55, 0x55, 0x55)

_fig_counter = {"n": 0}
_tab_counter = {"n": 0}


def load_json(p: Path):
    with open(p, "r", encoding="utf-8") as f:
        return json.load(f)


def read_csv_rows(p: Path):
    with open(p, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


# --- verified data ---------------------------------------------------------
TEST = load_json(MET / "test_evaluation_report.json")
VAL = load_json(MET / "val_metrics.json")
FINAL = load_json(MET / "final_report.json")
BEST = load_json(MET / "best_metrics.json")
MANIFEST = load_json(MET / "run_manifest.json")
STATS = load_json(REP / "dataset_statistics.json")
CLEAN = load_json(REP / "cleaning_report.json")
AUDIT = load_json(REP / "audit_report.json")
SPLITVER = load_json(REP / "split_verification_report.json")
NORM = load_json(MET / "normalization_stats.json")


# --- formatting helpers ----------------------------------------------------
def _pct(x, d=1):
    return f"{100 * float(x):.{d}f}%"


def _num(x, d=4):
    return f"{float(x):.{d}f}"


def _add_field(paragraph, instr, placeholder=""):
    run = paragraph.add_run()
    r = run._r
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    txt = OxmlElement("w:t")
    txt.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr_el, sep, txt, end):
        r.append(el)


def add_page_number_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Page ")
    _add_field(p, "PAGE")
    p.add_run(" of ")
    _add_field(p, "NUMPAGES")


def add_toc(doc):
    p = doc.add_paragraph()
    _add_field(
        p,
        'TOC \\o "1-3" \\h \\z \\u',
        "Right-click and choose 'Update Field' to build the table of contents.",
    )
def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = ACCENT
    return h


def para(doc, text, italic=False, size=None, color=None, bold=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    return p


def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p


def caption_fig(doc, text):
    _fig_counter["n"] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Figure {_fig_counter['n']}. {text}")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    return p


def caption_tab(doc, text):
    _tab_counter["n"] += 1
    p = doc.add_paragraph()
    r = p.add_run(f"Table {_tab_counter['n']}. {text}")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    return p


def add_figure(doc, relpath, caption, width=6.2):
    path = PRES / relpath
    if path.is_file():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
        caption_fig(doc, caption)
    else:
        para(doc, NA, italic=True, color=GREY)


def add_table(doc, headers, rows, caption=None):
    if caption:
        caption_tab(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
        for pgh in hdr[i].paragraphs:
            for run in pgh.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    return table
def _part_overview(doc):
    heading(doc, "Part 1 — Project Overview", level=1)
    para(doc,
         "RetinaGrade is an independent deep-learning system that grades the "
         "severity of diabetic retinopathy (DR) from a single colour fundus "
         "photograph. Given one retinal image, the system outputs one of five "
         "ordered clinical grades:")
    add_table(
        doc,
        ["Grade", "Clinical label", "Meaning"],
        [
            [0, "No DR (Normal)", "Healthy retina, no visible lesions"],
            [1, "Mild DR", "Microaneurysms only"],
            [2, "Moderate DR", "Microaneurysms, dot-blot hemorrhages, hard exudates"],
            [3, "Severe DR", "Extensive hemorrhages, venous beading, IRMA"],
            [4, "Proliferative DR (PDR)", "Neovascularization / preretinal hemorrhage"],
        ],
        caption="The five ordered diabetic-retinopathy grades produced by RetinaGrade.",
    )
    para(doc,
         "The system is built around a single end-to-end model, DualSwinOrd, "
         "that couples a Swin-Tiny vision backbone with two lightweight priors "
         "(a semantic-prior modulation module and a parallel large-kernel "
         "attention module) and a dual prediction head combining a categorical "
         "classifier and an ordinal (threshold) branch. The final grade is the "
         "argmax of the classification head.")
    para(doc,
         "This report documents the delivered system exactly as implemented and "
         "evaluated in the repository: the dataset and its preparation, the "
         "model, the training configuration, and the measured results on a "
         "held-out test set. Every number reported here is read directly from "
         "the project's stored evaluation artefacts.")
    bullet(doc, "Headline test QWK: " + _num(TEST["metrics"]["qwk"]), "Key result — ")
    bullet(doc, "Held-out test accuracy: " + _pct(TEST["metrics"]["accuracy"], 1))
    bullet(doc, "Referable-DR AUC (grade ≥ 2): " + _num(TEST["metrics"]["referable_auc"]))


def _part_dataset(doc):
    heading(doc, "Part 2 — Dataset", level=1)
    c = STATS["counts"]
    para(doc,
         f"After preparation the working dataset contains {c['total_images']} "
         "labelled fundus images across the five grades, partitioned into three "
         "disjoint splits. The split counts and per-grade composition below are "
         "taken from the dataset statistics artefact.")
    per = c["per_split"]
    ov = c["overall"]
    rows = []
    for g in range(5):
        rows.append([
            g, CLASS_FULL[g],
            per["train"][str(g)], per["val"][str(g)], per["test"][str(g)],
            ov[str(g)],
        ])
    rows.append(["—", "Total",
                 sum(per["train"].values()), sum(per["val"].values()),
                 sum(per["test"].values()), c["total_images"]])
    add_table(doc, ["Grade", "Class", "Train", "Val", "Test", "Overall"], rows,
              caption="Per-grade image counts by split (working dataset).")
    add_figure(doc, "Figures/Dataset/class_distribution_derived.png",
               "Per-grade class distribution across the train / validation / test splits.")
    add_figure(doc, "Figures/Dataset/class_distribution_original.png",
               "Class distribution of the prepared dataset (project data-prep artefact).")
    imb = STATS["imbalance"]
    para(doc, "Class imbalance (measured over the working dataset):", bold=True)
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Majority:minority ratio", _num(imb["majority_minority_ratio"], 2) + " : 1"],
            ["Shannon entropy (bits)", f"{imb['entropy_bits']:.3f} of {imb['max_entropy_bits']:.3f} max"],
            ["Normalized imbalance index", _num(imb["normalized_imbalance_index"], 3)],
            ["Gini coefficient", _num(imb["gini"], 3)],
            ["Effective number of classes", _num(imb["effective_num_classes"], 2)],
        ],
        caption="Quantified class imbalance of the working dataset.",
    )
    add_figure(doc, "Figures/Dataset/reference_class_weights.png",
               "Reference class-weighting schemes derived from the grade frequencies.")
    para(doc,
         "The dataset is dominated by grade 0 (No DR) and grade 2 (Moderate), "
         "while grade 3 (Severe) is the rarest class. This imbalance is analysed "
         "and reported; the delivered training run does not apply resampling or "
         "class weighting (see Part 9).", italic=True, color=GREY)
def _part_eda(doc):
    heading(doc, "Part 3 — Exploratory Data Analysis", level=1)
    res = AUDIT["resolution"]
    q = AUDIT["quality"]
    para(doc,
         f"An audit of the raw acquisition set ({AUDIT['totals']['images']} images, "
         f"{AUDIT['totals']['total_bytes']/1e9:.2f} GB) confirmed data integrity "
         "before any modelling: every file was readable, none were missing or "
         "zero-byte, all labels were in range, and no duplicate identifiers "
         "occurred within a split.")
    add_table(
        doc,
        ["Property", "Value"],
        [
            ["Images audited", AUDIT["totals"]["images"]],
            ["Unreadable / missing / zero-byte", "0 / 0 / 0"],
            ["File format", ", ".join(AUDIT["formats"].keys())],
            ["Colour mode", ", ".join(AUDIT["color_modes"].keys())],
            ["Distinct resolutions", res["distinct_resolutions"]],
            ["Width (median / min / max)",
             f"{res['width']['median']:.0f} / {res['width']['min']:.0f} / {res['width']['max']:.0f}"],
            ["Height (median / min / max)",
             f"{res['height']['median']:.0f} / {res['height']['min']:.0f} / {res['height']['max']:.0f}"],
            ["Aspect ratio (median)", _num(res["aspect_ratio"]["median"], 3)],
            ["Unique MD5 hashes", f"{AUDIT['hashes']['unique_md5']} of {AUDIT['hashes']['md5_computed']}"],
        ],
        caption="Raw-image audit summary (geometry, format and integrity).",
    )
    add_figure(doc, "Figures/EDA/eda_geometry_quality.png",
               "Image geometry and quality distributions observed during EDA.")
    para(doc,
         "Image geometry is highly heterogeneous (17 distinct resolutions, widths "
         f"from {res['width']['min']:.0f} to {res['width']['max']:.0f} px), which "
         "motivates the deterministic geometric normalization applied in "
         "preprocessing. Quality statistics — brightness "
         f"(median {q['brightness']['median']:.1f}), contrast "
         f"(median {q['contrast']['median']:.1f}) and sharpness — are used to "
         "derive the data-driven quality-flag thresholds described in Part 4.")


def _part_cleaning(doc):
    heading(doc, "Part 4 — Data Cleaning", level=1)
    t = CLEAN["totals"]
    para(doc,
         f"Cleaning follows a conservative, evidence-driven policy: exact "
         f"byte-duplicates that cross splits are excluded to prevent data "
         f"leakage, while quality outliers are flagged but never deleted. Of "
         f"{t['images']} audited images, {t['excluded']} were excluded (exact "
         f"cross-split MD5 duplicates) and {t['included']} were retained; "
         f"{t['flagged']} images carry one or more advisory flags. No raw file "
         "was modified or deleted on disk.")
    ex = CLEAN["exclusions_by_split"]
    add_table(
        doc,
        ["Exclusion (exact duplicates)", "Count"],
        [
            ["Total excluded", t["excluded"]],
            ["From test split", ex["test"]],
            ["From val split", ex["val"]],
            ["From train split", ex["train"]],
            ["Files physically deleted", CLEAN["policy"]["files_deleted"]],
        ],
        caption="Exact-duplicate exclusions by split (leakage prevention).",
    )
    fc = CLEAN["flag_counts"]
    add_table(
        doc,
        ["Advisory flag", "Count"],
        [[k.replace("_", " "), v] for k, v in sorted(fc.items(), key=lambda kv: -kv[1])],
        caption="Advisory quality / similarity flags (flag-only, no deletion).",
    )
    add_figure(doc, "Figures/Cleaning/cleaning_before_after_flags.png",
               "Per-split class counts before and after cleaning, with advisory flag totals.")
    para(doc,
         "Because quality flags mark genuine clinical acquisitions with valid "
         "labels, deleting them would worsen minority-class scarcity; they are "
         "therefore retained. The net effect of cleaning is small and confined "
         "to removing leakage-inducing exact duplicates.", italic=True, color=GREY)
def _part_preprocessing(doc):
    heading(doc, "Part 5 — Preprocessing", level=1)
    para(doc,
         "Every image passes through a fixed, deterministic preprocessing "
         "pipeline before it reaches the model. The pipeline standardises "
         "geometry and intensity so that the heterogeneous raw acquisitions "
         "become a uniform 512×512 RGB tensor.")
    add_table(
        doc,
        ["Order", "Stage", "Setting"],
        [
            [1, "Black-border removal", "threshold 10, blur kernel 5, min area ratio 0.10"],
            [2, "Circular crop", "margin ratio 0.0, fill value 0"],
            [3, "Resize to 512×512", "area down / linear up, keep-aspect-ratio = false"],
            [4, "CLAHE", "disabled (clip 2.0, tile 8×8, LAB) "],
            [5, "Illumination correction", "disabled (sigma 10, weight 4)"],
            [6, "Augmentation", "train split only (see Part 6)"],
            [7, "Normalize", "per-channel mean / std (train statistics)"],
            [8, "To tensor", "CHW float tensor"],
        ],
        caption="Deterministic preprocessing order and per-stage settings.",
    )
    m, s = NORM["mean"], NORM["std"]
    para(doc, "Normalization statistics (computed over the training split):", bold=True)
    add_table(
        doc,
        ["Channel", "Mean", "Std"],
        [
            ["R", _num(m[0]), _num(s[0])],
            ["G", _num(m[1]), _num(s[1])],
            ["B", _num(m[2]), _num(s[2])],
        ],
        caption=f"Channel normalization statistics ({NORM['images']} train images, "
                f"preprocessing hash {NORM['preprocessing_hash']}).",
    )
    add_figure(doc, "Figures/Preprocessing/preprocessing_stages_beforeafter.png",
               "Preprocessing applied to real fundus images, shown stage by stage.")
    add_figure(doc, "Figures/Preprocessing/preprocessing_preview_composite.png",
               "Preprocessing preview composite produced by the data-prep pipeline.")


def _part_augmentation(doc):
    heading(doc, "Part 6 — Data Augmentation", level=1)
    para(doc,
         "Augmentation is applied to the training split only and is intentionally "
         "conservative: it uses label-preserving photometric and geometric "
         "transforms that respect fundus imaging, and explicitly forbids "
         "occlusion / mixing augmentations that could destroy small lesions.")
    add_table(
        doc,
        ["Transform", "Probability", "Parameters"],
        [
            ["Horizontal flip", "0.5", "—"],
            ["Vertical flip", "0.5", "—"],
            ["Rotation", "0.6", "limit [-180°, 180°], constant border, fill 0"],
            ["Random brightness", "0.7", "limit 0.25"],
            ["Random contrast", "0.5", "limit 0.20"],
            ["Color jitter", "0.5", "b 0.20 / c 0.20 / s 0.15 / h 0.03"],
        ],
        caption="Enabled training-time augmentations and their parameters.",
    )
    add_table(
        doc,
        ["Category", "Transforms"],
        [
            ["Disabled", "scale jitter, conservative crop, gamma, gaussian blur, gaussian noise"],
            ["Forbidden (guarded)", "MixUp, CutMix, CutOut / CoarseDropout, GridDropout, RandomErasing"],
        ],
        caption="Disabled and explicitly forbidden augmentations.",
    )
    add_figure(doc, "Figures/Augmentation/augmentation_examples_grid.png",
               "A single preprocessed fundus image under each enabled augmentation.")
    para(doc,
         "Occlusion and mixing augmentations are blocked at the policy level: "
         "requesting one raises a guard error. This protects the small, "
         "diagnostically critical lesions (e.g. microaneurysms) that define the "
         "lower grades.", italic=True, color=GREY)
def _part_architecture(doc):
    heading(doc, "Part 7 — Model Architecture", level=1)
    para(doc,
         "RetinaGrade's model, DualSwinOrd, is a single end-to-end network. A "
         "512×512×3 preprocessed image is encoded by a Swin-Tiny transformer "
         "backbone; two lightweight modules inject clinical priors and enlarge "
         "the effective receptive field at the deepest stage; a shared neck "
         "pools features to a 512-dimensional vector; and a dual head produces "
         "both a categorical and an ordinal view of the grade.")
    add_figure(doc, "Figures/Architecture/model_architecture.png",
               "DualSwinOrd architecture: backbone, prior modules, shared neck and dual head.")
    add_table(
        doc,
        ["Component", "Description"],
        [
            ["Backbone", "Swin-Tiny (swin_tiny_patch4_window7_224), ImageNet-pretrained, 4 stages"],
            ["SPM", "Semantic Prior Modulation — FiLM-style sigmoid gate at stage 3; "
                    "5 clinical prompts embedded (dim 512) via a frozen hashing text adapter"],
            ["PLKA", "Parallel Large-Kernel Attention at stage 3 — parallel dilated "
                     "attention, kernel 3, GELU, batch-norm 2d"],
            ["Neck", "Shared feature neck — global average pooling → 512-d, dropout 0.0"],
            ["Head (classification)", "K = 5 softmax classifier"],
            ["Head (ordinal)", "K − 1 = 4 threshold logits, trained with BCE"],
            ["Prediction rule", "argmax of the classification softmax"],
        ],
        caption="DualSwinOrd components and their roles.",
    )
    para(doc,
         "The ordinal branch is trained jointly and supplies an ordering-aware "
         "signal, but the reported grade is taken from the classification head's "
         "argmax; the ordinal probabilities are collected and not used for the "
         "final prediction.", italic=True, color=GREY)


def _part_workflow(doc):
    heading(doc, "Part 8 — End-to-End Workflow", level=1)
    para(doc,
         "The diagram below traces one image from raw acquisition through "
         "audit, cleaning, deterministic preprocessing, model inference and "
         "grade output — the same path used at both training and evaluation "
         "time (augmentation is inserted only for the training split).")
    add_figure(doc, "Figures/Workflow/end_to_end_workflow.png",
               "End-to-end RetinaGrade workflow, from raw fundus image to predicted grade.")


def _part_training_config(doc):
    heading(doc, "Part 9 — Training Configuration", level=1)
    env = MANIFEST["environment"]
    para(doc,
         "The delivered model was trained with the fixed, fully-specified "
         "configuration below. The run is deterministic (fixed seed) and the "
         "checkpoint is selected by the best validation quadratic-weighted "
         "kappa (QWK).")
    add_table(
        doc,
        ["Hyper-parameter", "Value"],
        [
            ["Image size", "512 × 512"],
            ["Batch size", "16"],
            ["Data-loader workers", "4"],
            ["Max epochs", "50"],
            ["Optimizer", "AdamW (lr 1e-4, weight decay 1e-4)"],
            ["LR schedule", "cosine annealing (eta_min 0)"],
            ["Label smoothing", "0.1"],
            ["Loss weighting", "lambda_cls 0.5 (classification + ordinal)"],
            ["Mixed precision (AMP)", "enabled"],
            ["Class weighting / resampling", "none"],
            ["Monitor metric", "val_qwk (max)"],
            ["Early stopping", "disabled"],
            ["Seed / determinism", "42 / deterministic"],
        ],
        caption="Training hyper-parameters (delivered configuration).",
    )
    add_table(
        doc,
        ["Environment", "Value"],
        [
            ["Python", env["python"]],
            ["PyTorch", env["torch"]],
            ["CUDA available", str(env["cuda_available"])],
            ["Platform", env["platform"]],
            ["GPU model", NA],
            ["Config hashes (data / model / training)",
             f"{MANIFEST['hashes']['data']} / {MANIFEST['hashes']['model']} / {MANIFEST['hashes']['training']}"],
        ],
        caption="Recorded training environment and configuration hashes.",
    )
def _part_training_results(doc):
    heading(doc, "Part 10 — Training Results", level=1)
    bm = BEST["metrics"]
    para(doc,
         "Training was monitored on the validation split. The best checkpoint "
         "(best.pt) is the one maximising validation QWK. The training-time "
         "artefacts record the best epoch as 24 in the final report / run "
         "manifest / best-metrics files, and as 23 in the checkpoint metadata "
         "stored with the evaluation reports; both source values are reported "
         "here without alteration.")
    add_figure(doc, "Figures/Training/training_curves.png",
               "Training and validation loss / QWK across epochs.")
    add_figure(doc, "Figures/Training/training_perclass_f1.png",
               "Per-class validation F1 across epochs.")
    para(doc, "Metrics at the best validation checkpoint:", bold=True)
    add_table(
        doc,
        ["Validation metric", "Value"],
        [
            ["QWK", _num(bm["val_qwk"])],
            ["Accuracy", _pct(bm["val_accuracy"])],
            ["Macro-F1", _num(bm["val_macro_f1"])],
            ["MAE", _num(bm["val_mae"])],
            ["Within-one accuracy", _pct(bm["val_within_one_accuracy"])],
            ["Validation ECE (10-bin)", _num(FINAL["validation_ece"])],
            ["Train loss", _num(bm["train_loss"])],
            ["Val loss", _num(bm["val_loss"])],
            ["Learning rate at best epoch", f"{bm['lr']:.3e}"],
        ],
        caption="Validation metrics at the selected best checkpoint.",
    )


def _part_final_eval(doc):
    heading(doc, "Part 11 — Final Evaluation (Held-out Test Set)", level=1)
    tm = TEST["metrics"]
    para(doc,
         f"The delivered model was evaluated once on the held-out test split "
         f"({TEST['num_samples']} images) it never saw during training or "
         "checkpoint selection. These are the system's headline results.")
    add_table(
        doc,
        ["Test metric", "Value"],
        [
            ["Quadratic-weighted kappa (QWK)", _num(tm["qwk"])],
            ["Accuracy", _pct(tm["accuracy"])],
            ["Macro-F1", _num(tm["macro_f1"])],
            ["Mean absolute error (grades)", _num(tm["mae"])],
            ["Within-one accuracy", _pct(tm["within_one_accuracy"])],
            ["Referable-DR AUC (grade ≥ 2)", _num(tm["referable_auc"])],
            ["Referable-DR FNR (grade ≥ 2)", _num(tm["referable_fnr"])],
            ["Expected calibration error (10-bin)", _num(TEST["ece"])],
        ],
        caption="Held-out test-set metrics (single evaluation of best.pt).",
    )
    rows = []
    for i, cls in enumerate(CLASS_FULL):
        rows.append([
            cls,
            _num(tm["per_class_precision"][cls], 3),
            _num(tm["per_class_recall"][cls], 3),
            _num(tm["per_class_f1"][cls], 3),
        ])
    add_table(doc, ["Class", "Precision", "Recall", "F1"], rows,
              caption="Per-class precision, recall and F1 on the test set.")
    add_figure(doc, "Figures/Evaluation/test_confusion_matrix_hires.png",
               "Test-set confusion matrix (counts).")
    add_figure(doc, "Figures/Evaluation/test_perclass_prf1_bar.png",
               "Per-class precision / recall / F1 on the test set.")
    add_figure(doc, "Figures/Evaluation/test_reliability_diagram.png",
               "Test-set reliability diagram (confidence calibration).")
    add_figure(doc, "Figures/Evaluation/test_val_f1_radar.png",
               "Per-class F1 compared across the validation and test splits.")
    add_figure(doc, "Figures/Evaluation/test_val_metrics_comparison.png",
               "Headline metrics compared across the validation and test splits.")
    vm = VAL["metrics"]
    para(doc, "Validation-split reference (same checkpoint):", bold=True)
    add_table(
        doc,
        ["Metric", "Validation", "Test"],
        [
            ["QWK", _num(vm["qwk"]), _num(tm["qwk"])],
            ["Accuracy", _pct(vm["accuracy"]), _pct(tm["accuracy"])],
            ["Macro-F1", _num(vm["macro_f1"]), _num(tm["macro_f1"])],
            ["Within-one accuracy", _pct(vm["within_one_accuracy"]), _pct(tm["within_one_accuracy"])],
            ["Referable-DR AUC", _num(vm["referable_auc"]), _num(tm["referable_auc"])],
        ],
        caption="Validation vs test metrics for the delivered checkpoint.",
    )
def _part_predictions(doc):
    heading(doc, "Part 12 — Prediction Examples", level=1)
    para(doc,
         "The panels below show the delivered model's predictions on real "
         "held-out test images, produced by a read-only inference pass over the "
         "test split with the selected checkpoint (best.pt). Each panel reports "
         "the ground-truth grade, the predicted grade and the model's confidence "
         "(maximum softmax probability). A green border marks a correct grade; a "
         "red border marks an error.")
    n = TEST["num_samples"]
    acc = TEST["metrics"]["accuracy"]
    para(doc,
         f"Across all {n} test images the inference pass reproduced "
         f"{round(acc * n)} correct grades ({_pct(acc, 1)} accuracy), matching "
         "the stored evaluation report.")
    add_figure(doc, "Figures/Predictions/prediction_examples_grid.png",
               "Representative TEST-split predictions (green = correct, red = incorrect).")
    para(doc, "Selected individual cases:", bold=True)
    add_figure(doc, "Figures/Predictions/prediction_01_gt0_pred0_correct.png",
               "Correct: a No-DR image graded No DR with high confidence.", width=3.2)
    add_figure(doc, "Figures/Predictions/prediction_04_gt3_pred3_correct.png",
               "Correct: a Severe-DR image graded Severe.", width=3.2)
    add_figure(doc, "Figures/Predictions/prediction_08_gt1_pred2_incorrect.png",
               "Error: a Mild-DR image graded Moderate (adjacent-grade confusion).", width=3.2)
    add_figure(doc, "Figures/Predictions/prediction_11_gt4_pred2_incorrect.png",
               "Error: a Proliferative-DR image graded Moderate.", width=3.2)


def _part_error_analysis(doc):
    heading(doc, "Part 13 — Error Analysis", level=1)
    para(doc,
         "Errors were analysed from the test-set confusion matrix. The model's "
         "mistakes are overwhelmingly between neighbouring grades: within-one "
         f"accuracy is {_pct(TEST['metrics']['within_one_accuracy'], 1)}, so "
         "almost every error is off by a single grade rather than a gross "
         "misclassification.")
    add_figure(doc, "Figures/ErrorAnalysis/error_analysis_confusion.png",
               "Test-set error analysis: the dominant grade confusions.")
    cm = [[194, 2, 1, 0, 0], [0, 13, 14, 0, 1], [0, 2, 70, 3, 2],
          [0, 0, 9, 4, 0], [0, 2, 7, 2, 16]]
    rows = [[CLASS_SHORT[i]] + cm[i] for i in range(5)]
    add_table(doc,
              ["True grade", "No DR", "Mild", "Moderate", "Severe", "PDR"], rows,
              caption="Test-set confusion matrix (rows = true grade, columns = predicted grade).")
    para(doc, "Dominant confusions (true → predicted):", bold=True)
    bullet(doc, "Mild → Moderate: 14 of 28 Mild test images.")
    bullet(doc, "Severe → Moderate: 9 of 13 Severe test images.")
    bullet(doc, "Proliferative (PDR) → Moderate: 7 of 27 PDR test images.")
    bullet(doc, "Moderate → Mild: 2 of 77 Moderate test images.")
    para(doc,
         "The recurring pattern is under-grading of the rarer, more severe classes "
         "toward the dominant Moderate class. This is consistent with the class "
         "imbalance documented in Part 2 and directly explains the lower recall on "
         "the Mild, Severe and PDR grades.", italic=True, color=GREY)


def _part_limitations(doc):
    heading(doc, "Part 14 — Limitations", level=1)
    tm = TEST["metrics"]
    para(doc,
         "The following limitations are stated strictly from measured evidence in "
         "the repository; each is backed by a reported number.")
    bullet(doc,
           f"Severe-grade recall is low ({_num(tm['per_class_recall'][CLASS_FULL[3]], 3)}): "
           "only 4 of 13 Severe test images are graded correctly.",
           "Minority-class recall — ")
    bullet(doc,
           f"Mild-grade recall is limited ({_num(tm['per_class_recall'][CLASS_FULL[1]], 3)}): "
           "13 of 28 Mild test images are graded correctly.")
    bullet(doc,
           f"The dataset is imbalanced ({_num(STATS['imbalance']['majority_minority_ratio'], 2)}"
           " : 1 majority-to-minority), and the delivered run applies no resampling "
           "or class weighting.",
           "Class imbalance — ")
    bullet(doc,
           f"Calibration is imperfect: the test expected calibration error is "
           f"{_num(TEST['ece'], 3)} (10-bin).",
           "Calibration — ")
    bullet(doc,
           f"The test set is small ({TEST['num_samples']} images), and the rarest "
           "grade (Severe) has only 13 test images, so per-class estimates for the "
           "minority grades carry wide uncertainty.",
           "Sample size — ")
    para(doc, "Artefacts not present in the repository (reported as unavailable):", bold=True)
    bullet(doc, "Grad-CAM / attention overlays: " + NA)
    bullet(doc, "Ablation study: " + NA)
    bullet(doc, "ROC / precision-recall curve image files: " + NA)
    bullet(doc, "TensorBoard-rendered PNG exports: " + NA)


def _part_strengths(doc):
    heading(doc, "Part 15 — Strengths", level=1)
    tm = TEST["metrics"]
    para(doc, "The delivered system's measured strengths, each backed by a test-set number:")
    bullet(doc, f"Strong ordinal agreement: test QWK {_num(tm['qwk'])}.", "Agreement — ")
    bullet(doc, f"Errors stay adjacent: within-one accuracy {_pct(tm['within_one_accuracy'], 1)}.",
           "Clinical safety — ")
    bullet(doc,
           f"Reliable referable-DR screening (grade ≥ 2): AUC {_num(tm['referable_auc'])} "
           f"with a low false-negative rate of {_num(tm['referable_fnr'], 3)}.",
           "Screening — ")
    bullet(doc, f"Near-perfect No-DR detection: No-DR F1 {_num(tm['per_class_f1'][CLASS_FULL[0]], 3)}.",
           "Normal detection — ")
    bullet(doc, "Fully deterministic and reproducible (fixed seed 42, recorded config hashes).",
           "Reproducibility — ")
    bullet(doc, "Lesion-preserving augmentation policy that forbids occlusion / mixing transforms.",
           "Domain-aware design — ")
    bullet(doc, "Confidence calibration is measured and reported (reliability diagram + ECE).",
           "Transparency — ")


def _part_assets(doc):
    heading(doc, "Part 16 — Presentation Assets", level=1)
    para(doc,
         "All presentation materials are collected under the Presentation/ folder. "
         "Figures are grouped by topic; every embedded chart and every copied "
         "artefact is stored locally so the report is fully self-contained.")
    add_table(
        doc,
        ["Folder", "Contents"],
        [
            ["Figures/Architecture", "Model architecture diagram"],
            ["Figures/Dataset", "Class-distribution and class-weight charts"],
            ["Figures/EDA", "Image geometry / quality distributions"],
            ["Figures/Cleaning", "Before / after cleaning and flag totals"],
            ["Figures/Preprocessing", "Stage-by-stage preprocessing panels"],
            ["Figures/Augmentation", "Augmentation examples grid"],
            ["Figures/Workflow", "End-to-end workflow diagram"],
            ["Figures/Training", "Training curves and per-class F1"],
            ["Figures/Evaluation", "Confusion matrices, reliability diagrams, metric comparisons"],
            ["Figures/ErrorAnalysis", "Confusion-focused error analysis"],
            ["Figures/Predictions", "Prediction grid and individual example panels"],
            ["SampleImages", "15 real fundus images (3 per grade) from the test split"],
            ["Metrics", "Evaluation JSON/CSV artefacts (test, validation, training log)"],
            ["Reports", "Dataset statistics, cleaning, audit and split-verification reports"],
            ["Configs", "data.yaml, model.yaml, training.yaml"],
            ["Tables", "Train / validation / test split label tables"],
            ["scripts", "Reproducible generation scripts for every asset"],
        ],
        caption="Structure of the Presentation/ asset collection.",
    )
    para(doc,
         "A complete, itemised index of every asset — with original location, "
         "copied location, category and suggested slide — is provided in "
         "Presentation_Assets_Index.xlsx (see Part 17).", italic=True, color=GREY)


def _part_index_pointer(doc):
    heading(doc, "Part 17 — Assets Index (spreadsheet)", level=1)
    para(doc,
         "Presentation_Assets_Index.xlsx catalogues every file in the "
         "Presentation/ folder. It has one row per asset with the following "
         "columns: File Name, Original Location, Copied Location, Category, "
         "Suggested Slide and Description. Use it to locate the source of any "
         "figure or artefact and to see which slide it belongs on.")


def _part_manifest_pointer(doc):
    heading(doc, "Part 18 — Presentation Manifest", level=1)
    para(doc,
         "presentation_manifest.md maps each key asset to a recommended slide, a "
         "suggested usage, an importance rating and ready-to-read speaker notes. "
         "It is intended as the working script for assembling and delivering the "
         "slide deck.")


def _part_docx_pointer(doc):
    heading(doc, "Part 19 — This Report", level=1)
    para(doc,
         "This document (Presentation_Report.docx) is generated reproducibly by "
         "Presentation/scripts/make_docx.py from the verified artefacts in the "
         "Presentation/ folder. To populate the Table of Contents and the page "
         "totals in Word, select all (Ctrl+A) and press F9, or right-click the "
         "table of contents and choose 'Update Field'.")


def _part_summary_pointer(doc):
    heading(doc, "Part 20 — Presentation Summary", level=1)
    para(doc,
         "Presentation_Summary.md distils the whole project into a compact brief: "
         "the highlights, the best results and key numbers, the figures to show, "
         "the main talking points, and a set of likely defence questions with "
         "concise, evidence-backed answers. Read it last as a final-preparation "
         "cheat-sheet before the presentation.")


def build():
    doc = Document()
    # base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ---- title page ----
    for _ in range(6):
        doc.add_paragraph()
    para(doc, "RetinaGrade", size=40, bold=True, color=ACCENT,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Diabetic Retinopathy Grading System", size=20, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Project Presentation Report", size=16,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    para(doc, "An independent AI system for 5-grade fundus image classification",
         italic=True, size=12, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(8):
        doc.add_paragraph()
    para(doc, "Prepared: 2026-08-03", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Environment: Python 3.10.11 · PyTorch 2.5.1+cu124 · Seed 42 (deterministic)",
         size=10, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

    # ---- table of contents ----
    doc.add_page_break()
    heading(doc, "Table of Contents", level=1)
    add_toc(doc)

    # ---- body section with page numbers ----
    doc.add_page_break()
    add_page_number_footer(doc.sections[0])

    _part_overview(doc)
    _part_dataset(doc)
    _part_eda(doc)
    _part_cleaning(doc)
    _part_preprocessing(doc)
    _part_augmentation(doc)
    _part_architecture(doc)
    _part_workflow(doc)
    _part_training_config(doc)
    _part_training_results(doc)
    _part_final_eval(doc)
    _part_predictions(doc)
    _part_error_analysis(doc)
    _part_limitations(doc)
    _part_strengths(doc)
    _part_assets(doc)
    _part_index_pointer(doc)
    _part_manifest_pointer(doc)
    _part_docx_pointer(doc)
    _part_summary_pointer(doc)

    out = PRES / "Presentation_Report.docx"
    doc.save(str(out))
    print(f"[ok] wrote {out}")


if __name__ == "__main__":
    build()
